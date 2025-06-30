from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Query
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime
import io
import PyPDF2
import httpx
import json
from emergentintegrations.llm.chat import LlmChat, UserMessage
import psutil
import subprocess
import sys
import pkg_resources
from typing import Union
import asyncio
import time
import threading

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Environment configuration with defaults for local development
MONGO_URL = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
DB_NAME = os.environ.get('DB_NAME', 'chatpdf_database')

# Load multiple OpenRouter API keys for load balancing and fallback
OPENROUTER_API_KEYS = []
for i in range(1, 6):  # Load up to 5 keys
    key_name = 'OPENROUTER_API_KEY' if i == 1 else f'OPENROUTER_API_KEY_{i}'
    key_value = os.environ.get(key_name, '')
    if key_value:
        OPENROUTER_API_KEYS.append(key_value)

# Backward compatibility - set primary key
OPENROUTER_API_KEY = OPENROUTER_API_KEYS[0] if OPENROUTER_API_KEYS else ''

GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '')
ENVIRONMENT = os.environ.get('ENVIRONMENT', 'development')

# Add a counter for round-robin load balancing
import threading
_openrouter_key_counter = threading.local()

def get_next_openrouter_key():
    """Get the next OpenRouter API key using round-robin load balancing"""
    if not OPENROUTER_API_KEYS:
        return ''
    
    if not hasattr(_openrouter_key_counter, 'value'):
        _openrouter_key_counter.value = 0
    
    key = OPENROUTER_API_KEYS[_openrouter_key_counter.value % len(OPENROUTER_API_KEYS)]
    _openrouter_key_counter.value += 1
    return key

# Configure allowed origins based on environment
if ENVIRONMENT == 'production':
    ALLOWED_ORIGINS = [
        "https://*.netlify.app",
        "https://*.vercel.app", 
        "https://*.railway.app",
        os.environ.get('FRONTEND_URL', '')
    ]
    # Remove empty strings
    ALLOWED_ORIGINS = [origin for origin in ALLOWED_ORIGINS if origin]
else:
    ALLOWED_ORIGINS = [
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:3001"
    ]

# Validate that we have at least one AI provider configured
if not OPENROUTER_API_KEYS and not GEMINI_API_KEY:
    raise ValueError("At least one AI provider API key is required (OPENROUTER_API_KEY or GEMINI_API_KEY)")

# MongoDB connection
client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

# OpenRouter API configuration
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

# Create the main app
app = FastAPI(title="Baloch AI chat PdF & GPT API", version="2.0.0")
api_router = APIRouter(prefix="/api")

# AI Functions
def is_gemini_model(model: str) -> bool:
    """Check if the model is a Gemini model"""
    gemini_models = [
        'gemini-2.5-flash-preview-04-17',
        'gemini-2.5-pro-preview-05-06', 
        'gemini-2.0-flash',
        'gemini-2.0-flash-preview-image-generation',
        'gemini-2.0-flash-lite',
        'gemini-1.5-flash',
        'gemini-1.5-flash-8b',
        'gemini-1.5-pro'
    ]
    return model in gemini_models

async def get_ai_response_gemini(messages: List[Dict], model: str) -> str:
    """Handle Gemini API requests using emergentintegrations"""
    try:
        # Create a unique session ID for this conversation
        session_id = str(uuid.uuid4())
        
        # Extract system message
        system_message = next((msg["content"] for msg in messages if msg["role"] == "system"), "You are a helpful assistant.")
        
        # Initialize LlmChat with Gemini
        chat = LlmChat(
            api_key=GEMINI_API_KEY,
            session_id=session_id,
            system_message=system_message
        ).with_model("gemini", model)
        
        # Get the last user message (most recent)
        user_messages = [msg for msg in messages if msg["role"] == "user"]
        if not user_messages:
            raise HTTPException(status_code=400, detail="No user message found")
        
        last_user_message = user_messages[-1]["content"]
        
        # Create UserMessage and send
        user_message = UserMessage(text=last_user_message)
        response = await chat.send_message(user_message)
        
        return response
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gemini AI service error: {str(e)}")

async def get_ai_response_openrouter(messages: List[Dict], model: str) -> str:
    """Handle OpenRouter API requests (Claude models) with load balancing and fallback"""
    if not OPENROUTER_API_KEYS:
        raise HTTPException(status_code=500, detail="No OpenRouter API keys configured")
    
    # Convert chat format to OpenRouter format
    system_message = next((msg["content"] for msg in messages if msg["role"] == "system"), None)
    chat_messages = [{"role": msg["role"], "content": msg["content"]} for msg in messages if msg["role"] != "system"]
    
    # Try each API key with fallback logic
    last_error = None
    
    for attempt in range(len(OPENROUTER_API_KEYS)):
        # Get next key using round-robin
        api_key = get_next_openrouter_key()
        
        try:
            # Use OpenRouter API
            async with httpx.AsyncClient() as client:
                response = await client.post(
                    f"{OPENROUTER_BASE_URL}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "HTTP-Referer": "https://github.com/baloch/chatpdf",
                        "X-Title": "ChatPDF App"
                    },
                    json={
                        "model": model,
                        "messages": chat_messages,
                        "system": system_message,
                        "max_tokens": 2000,
                        "temperature": 0.7
                    }
                )
                response.raise_for_status()
                result = response.json()
                return result["choices"][0]["message"]["content"]
                
        except Exception as e:
            last_error = e
            logger.warning(f"OpenRouter API key {api_key[-10:]}... failed (attempt {attempt + 1}/{len(OPENROUTER_API_KEYS)}): {str(e)}")
            continue
    
    # If all keys failed, raise the last error
    raise HTTPException(status_code=500, detail=f"All OpenRouter API keys failed. Last error: {str(last_error)}")

async def get_ai_response(messages: List[Dict], model: str = "claude-3-opus-20240229") -> str:
    """Route AI requests to appropriate provider based on model"""
    try:
        if is_gemini_model(model):
            if not GEMINI_API_KEY:
                raise HTTPException(status_code=500, detail="Gemini API key not configured")
            return await get_ai_response_gemini(messages, model)
        else:
            if not OPENROUTER_API_KEYS:
                raise HTTPException(status_code=500, detail="OpenRouter API keys not configured")
            return await get_ai_response_openrouter(messages, model)
    except Exception as e:
        # If there's an error with the primary provider, try the backup
        if is_gemini_model(model) and OPENROUTER_API_KEYS:
            # If Gemini fails, try with a Claude model as backup
            logger.warning(f"Gemini model {model} failed, trying Claude backup: {str(e)}")
            return await get_ai_response_openrouter(messages, "claude-3-haiku-20240307")
        elif not is_gemini_model(model) and GEMINI_API_KEY:
            # If Claude fails, try with Gemini as backup
            logger.warning(f"Claude model {model} failed, trying Gemini backup: {str(e)}")
            return await get_ai_response_gemini(messages, "gemini-1.5-flash")
        else:
            raise e

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup_event():
    logger.info("🚀 Baloch AI chat PdF & GPT Backend starting up...")
    logger.info(f"📊 MongoDB URL: {MONGO_URL}")
    logger.info(f"🗄️  Database: {DB_NAME}")
    logger.info(f"🔑 OpenRouter API Keys: {'✅ ' + str(len(OPENROUTER_API_KEYS)) + ' keys configured' if OPENROUTER_API_KEYS else '❌ Missing'}")
    if OPENROUTER_API_KEYS:
        for i, key in enumerate(OPENROUTER_API_KEYS, 1):
            logger.info(f"   Key {i}: ...{key[-10:]}")
    logger.info(f"🤖 Gemini API Key: {'✅ Configured' if GEMINI_API_KEY else '❌ Missing'}")
    logger.info("✅ Baloch AI chat PdF & GPT Backend ready!")

@app.on_event("shutdown")
async def shutdown_db_client():
    logger.info("🛑 Shutting down Baloch AI chat PdF & GPT Backend...")
    client.close()
    logger.info("✅ Database connection closed")

# Health check models
class HealthMetrics(BaseModel):
    cpu_usage: float
    memory_usage: float
    disk_usage: float
    response_time: float
    active_sessions: int
    total_api_calls: int
    error_rate: float

class HealthIssue(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    issue_type: str  # 'critical', 'warning', 'performance'
    category: str  # 'service', 'api', 'database', 'dependency', 'performance'
    title: str
    description: str
    suggested_fix: str
    auto_fixable: bool
    severity: int  # 1-5 (5 being most critical)
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    resolved: bool = False
    resolution_time: Optional[datetime] = None

class SystemHealthStatus(BaseModel):
    overall_status: str  # 'healthy', 'warning', 'critical'
    backend_status: str
    frontend_status: str
    database_status: str
    api_status: str
    last_check: datetime = Field(default_factory=datetime.utcnow)
    metrics: HealthMetrics
    issues: List[HealthIssue] = []
    uptime: float  # seconds since last restart

class FixRequest(BaseModel):
    issue_id: str
    confirm_fix: bool = False

# Health check endpoint
@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow()}

# Comprehensive system health endpoint
@api_router.get("/system-health")
async def get_system_health():
    """Get comprehensive system health status"""
    try:
        health_status = await perform_comprehensive_health_check()
        return health_status
    except Exception as e:
        logger.error(f"Error performing health check: {e}")
        return SystemHealthStatus(
            overall_status="critical",
            backend_status="unhealthy",
            frontend_status="unknown",
            database_status="unknown",
            api_status="unknown",
            metrics=HealthMetrics(
                cpu_usage=0, memory_usage=0, disk_usage=0, response_time=0,
                active_sessions=0, total_api_calls=0, error_rate=100
            ),
            issues=[HealthIssue(
                issue_type="critical",
                category="service",
                title="Health Check Failed",
                description=str(e),
                suggested_fix="Investigate server logs and restart services",
                auto_fixable=True,
                severity=5
            )],
            uptime=0
        )

# Auto-fix endpoint with confirmation
@api_router.post("/system-health/fix")
async def fix_system_issue(request: FixRequest):
    """Apply auto-fix for system issues with user confirmation"""
    if not request.confirm_fix:
        return {"error": "Fix confirmation required", "confirmed": False}
    
    # Find the issue
    issue = None
    for health_issue in health_monitor_data["issues"]:
        if health_issue.id == request.issue_id:
            issue = health_issue
            break
    
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    
    if not issue.auto_fixable:
        raise HTTPException(status_code=400, detail="Issue is not auto-fixable")
    
    try:
        fix_result = await apply_auto_fix(issue)
        
        # Mark issue as resolved
        issue.resolved = True
        issue.resolution_time = datetime.utcnow()
        
        return {
            "success": True,
            "issue_id": request.issue_id,
            "fix_applied": fix_result["action"],
            "result": fix_result["result"],
            "message": fix_result["message"]
        }
    
    except Exception as e:
        logger.error(f"Error applying auto-fix: {e}")
        return {
            "success": False,
            "issue_id": request.issue_id,
            "error": str(e)
        }

async def apply_auto_fix(issue: HealthIssue) -> dict:
    """Apply automated fixes based on issue type"""
    
    if issue.category == "dependency":
        # Auto-install missing dependencies
        if "Missing packages:" in issue.description:
            # Extract package names from description
            packages = issue.description.split("Missing packages: ")[1].split(", ")
            install_results = []
            
            for package in packages:
                try:
                    result = subprocess.run([
                        sys.executable, "-m", "pip", "install", package
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0:
                        install_results.append(f"✅ {package} installed successfully")
                    else:
                        install_results.append(f"❌ {package} failed: {result.stderr}")
                except subprocess.TimeoutExpired:
                    install_results.append(f"⏱️ {package} installation timed out")
                except Exception as e:
                    install_results.append(f"❌ {package} error: {str(e)}")
            
            return {
                "action": "install_dependencies",
                "result": "partial" if any("❌" in r for r in install_results) else "success",
                "message": "\n".join(install_results)
            }
    
    elif issue.category == "performance":
        if "High CPU Usage" in issue.title or "High Memory Usage" in issue.title:
            # Clear system caches and suggest restart
            try:
                # Clear Python caches
                import gc
                gc.collect()
                
                # Clear any application caches if they exist
                # (You can add more cache clearing logic here)
                
                return {
                    "action": "clear_caches",
                    "result": "success",
                    "message": "System caches cleared. Consider restarting services if issue persists."
                }
            except Exception as e:
                return {
                    "action": "clear_caches",
                    "result": "failed",
                    "message": f"Cache clearing failed: {str(e)}"
                }
    
    elif issue.category == "service":
        # Service restart suggestions
        return {
            "action": "service_restart_suggestion",
            "result": "info",
            "message": "Consider restarting the backend service using: sudo supervisorctl restart backend"
        }
    
    elif issue.category == "database":
        # Database reconnection attempt
        try:
            global client, db
            client.close()
            client = AsyncIOMotorClient(MONGO_URL)
            db = client[DB_NAME]
            
            # Test connection
            await client.admin.command('ping')
            
            return {
                "action": "database_reconnect",
                "result": "success",
                "message": "Database connection re-established successfully"
            }
        except Exception as e:
            return {
                "action": "database_reconnect",
                "result": "failed",
                "message": f"Database reconnection failed: {str(e)}"
            }
    
    else:
        return {
            "action": "no_fix_available",
            "result": "info",
            "message": "No automated fix available for this issue type"
        }

# Get health metrics history
@api_router.get("/system-health/metrics")
async def get_health_metrics():
    """Get historical health metrics"""
    return {
        "current_metrics": get_system_metrics(),
        "history": health_monitor_data.get("metrics_history", [])[-50:],  # Last 50 data points
        "uptime": (datetime.utcnow() - health_monitor_data["start_time"]).total_seconds()
    }

# Middleware to track API calls and response times
@app.middleware("http")
async def track_api_metrics(request, call_next):
    start_time = time.time()
    
    try:
        response = await call_next(request)
        
        # Track successful API call
        health_monitor_data["api_calls"] += 1
        
        # Track response time
        response_time = (time.time() - start_time) * 1000  # milliseconds
        
        # Store metrics history (keep last 100 entries)
        metrics_entry = {
            "timestamp": datetime.utcnow().isoformat(),
            "response_time": response_time,
            "status_code": response.status_code,
            "endpoint": str(request.url.path)
        }
        
        health_monitor_data["metrics_history"].append(metrics_entry)
        if len(health_monitor_data["metrics_history"]) > 100:
            health_monitor_data["metrics_history"].pop(0)
        
        return response
        
    except Exception as e:
        # Track error
        health_monitor_data["errors"] += 1
        
        # Log the error
        logger.error(f"API Error: {str(e)} for {request.url.path}")
        
        raise e

# Pydantic Models
class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    content: str
    role: str  # 'user' or 'assistant'
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    feature_type: str = "chat"  # 'chat', 'qa_generation', 'general_ai', 'research'

class ChatSession(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)
    pdf_filename: Optional[str] = None
    pdf_content: Optional[str] = None

class PDFDocument(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    content: str
    upload_date: datetime = Field(default_factory=datetime.utcnow)
    file_size: int

class SendMessageRequest(BaseModel):
    session_id: str
    content: str
    model: str = "claude-3-opus-20240229"
    feature_type: str = "chat"

class CreateSessionRequest(BaseModel):
    title: str = "New Chat"

class GenerateQuestionsRequest(BaseModel):
    session_id: str
    question_type: str = "mixed"  # 'faq', 'mcq', 'true_false', 'mixed'
    chapter_segment: Optional[str] = None  # For chapter-specific questions
    model: str = "claude-3-opus-20240229"

class GenerateQuizRequest(BaseModel):
    session_id: str
    quiz_type: str = "daily"  # 'daily', 'manual'
    difficulty: str = "medium"  # 'easy', 'medium', 'hard'
    question_count: int = 10
    model: str = "claude-3-opus-20240229"


class TranslateRequest(BaseModel):
    session_id: str
    target_language: str
    content_type: str = "full"  # 'full', 'summary'
    model: str = "claude-3-opus-20240229"

class SearchRequest(BaseModel):
    query: str
    search_type: str = "all"  # 'all', 'pdfs', 'conversations'
    limit: int = 20

class ExportRequest(BaseModel):
    session_id: str
    export_format: str = "pdf"  # 'pdf', 'txt', 'docx'
    include_messages: bool = True
    feature_type: Optional[str] = None

# System Health Models
class HealthMetrics(BaseModel):
    cpu_usage: float
    memory_usage: float
    disk_usage: float
    response_time: float
    active_sessions: int
    total_api_calls: int
    error_rate: float

class HealthIssue(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    issue_type: str  # 'critical', 'warning', 'performance'
    category: str  # 'service', 'api', 'database', 'dependency', 'performance'
    title: str
    description: str
    suggested_fix: str
    auto_fixable: bool
    severity: int  # 1-5 (5 being most critical)
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    resolved: bool = False
    resolution_time: Optional[datetime] = None

class SystemHealthStatus(BaseModel):
    overall_status: str  # 'healthy', 'warning', 'critical'
    backend_status: str
    frontend_status: str
    database_status: str
    api_status: str
    last_check: datetime = Field(default_factory=datetime.utcnow)
    metrics: HealthMetrics
    issues: List[HealthIssue] = []
    uptime: float  # seconds since last restart

class FixRequest(BaseModel):
    issue_id: str
    confirm_fix: bool = False

# Global health monitoring variables
health_monitor_data = {
    "start_time": datetime.utcnow(),
    "api_calls": 0,
    "errors": 0,
    "issues": [],
    "last_health_check": None,
    "metrics_history": []
}

# Health Monitoring Functions
async def check_database_health() -> tuple[bool, str]:
    """Check MongoDB connection health"""
    try:
        # Try to ping the database
        await client.admin.command('ping')
        # Check if our database exists and is accessible
        collections = await db.list_collection_names()
        return True, f"Database healthy - {len(collections)} collections available"
    except Exception as e:
        return False, f"Database connection failed: {str(e)}"

async def check_api_keys_health() -> tuple[bool, str, dict]:
    """Check API keys validity"""
    api_status = {
        "openrouter": {"valid": [], "errors": []},
        "gemini": {"valid": False, "error": ""}
    }
    
    # Check OpenRouter API keys (all of them)
    if OPENROUTER_API_KEYS:
        for i, api_key in enumerate(OPENROUTER_API_KEYS, 1):
            try:
                async with httpx.AsyncClient() as client_http:
                    response = await client_http.get(
                        f"{OPENROUTER_BASE_URL}/models",
                        headers={"Authorization": f"Bearer {api_key}"},
                        timeout=10.0
                    )
                    if response.status_code == 200:
                        api_status["openrouter"]["valid"].append(f"Key {i}")
                    else:
                        api_status["openrouter"]["errors"].append(f"Key {i}: HTTP {response.status_code}")
            except Exception as e:
                api_status["openrouter"]["errors"].append(f"Key {i}: {str(e)}")
    else:
        api_status["openrouter"]["errors"].append("No API keys configured")
    
    # Check Gemini API (simplified check)
    if GEMINI_API_KEY:
        try:
            # Simple validation - just check if key is properly formatted
            if GEMINI_API_KEY.startswith("AIza") and len(GEMINI_API_KEY) > 30:
                api_status["gemini"]["valid"] = True
            else:
                api_status["gemini"]["error"] = "Invalid API key format"
        except Exception as e:
            api_status["gemini"]["error"] = str(e)
    else:
        api_status["gemini"]["error"] = "API key not configured"
    
    # Determine overall API health
    valid_openrouter_keys = len(api_status["openrouter"]["valid"])
    total_openrouter_keys = len(OPENROUTER_API_KEYS)
    gemini_valid = api_status["gemini"]["valid"]
    
    if valid_openrouter_keys == 0 and not gemini_valid:
        return False, "No valid API keys", api_status
    elif valid_openrouter_keys < total_openrouter_keys and gemini_valid:
        return True, f"OpenRouter: {valid_openrouter_keys}/{total_openrouter_keys} keys valid, Gemini: valid", api_status
    elif valid_openrouter_keys == total_openrouter_keys and gemini_valid:
        return True, "All API keys valid", api_status
    else:
        return True, f"OpenRouter: {valid_openrouter_keys}/{total_openrouter_keys} keys valid", api_status

async def check_dependencies() -> tuple[bool, str, list]:
    """Check if all required dependencies are installed"""
    required_packages = [
        'fastapi', 'uvicorn', 'motor', 'pymongo', 'httpx', 'PyPDF2',
        'anthropic', 'emergentintegrations', 'psutil', 'reportlab',
        'python-docx', 'python-dotenv', 'pydantic'
    ]
    
    missing_packages = []
    
    for package in required_packages:
        try:
            pkg_resources.get_distribution(package)
        except pkg_resources.DistributionNotFound:
            missing_packages.append(package)
    
    if missing_packages:
        return False, f"Missing packages: {', '.join(missing_packages)}", missing_packages
    else:
        return True, "All dependencies installed", []

def get_system_metrics() -> HealthMetrics:
    """Get current system performance metrics"""
    try:
        cpu_usage = psutil.cpu_percent()
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        # Calculate error rate
        total_calls = health_monitor_data["api_calls"]
        total_errors = health_monitor_data["errors"]
        error_rate = (total_errors / total_calls * 100) if total_calls > 0 else 0
        
        return HealthMetrics(
            cpu_usage=cpu_usage,
            memory_usage=memory.percent,
            disk_usage=disk.percent,
            response_time=0.0,  # Will be calculated during API calls
            active_sessions=0,  # Will be queried from database
            total_api_calls=total_calls,
            error_rate=error_rate
        )
    except Exception as e:
        logger.error(f"Error getting system metrics: {e}")
        return HealthMetrics(
            cpu_usage=0, memory_usage=0, disk_usage=0, response_time=0,
            active_sessions=0, total_api_calls=0, error_rate=0
        )

async def analyze_performance_issues(metrics: HealthMetrics) -> List[HealthIssue]:
    """Analyze metrics and identify performance issues"""
    issues = []
    
    # High CPU usage
    if metrics.cpu_usage > 80:
        issues.append(HealthIssue(
            issue_type="performance",
            category="performance",
            title="High CPU Usage",
            description=f"CPU usage is at {metrics.cpu_usage:.1f}%",
            suggested_fix="Consider restarting services or optimizing queries",
            auto_fixable=True,
            severity=3 if metrics.cpu_usage > 90 else 2
        ))
    
    # High memory usage
    if metrics.memory_usage > 80:
        issues.append(HealthIssue(
            issue_type="performance",
            category="performance",
            title="High Memory Usage",
            description=f"Memory usage is at {metrics.memory_usage:.1f}%",
            suggested_fix="Clear caches and restart services",
            auto_fixable=True,
            severity=3 if metrics.memory_usage > 90 else 2
        ))
    
    # High error rate
    if metrics.error_rate > 10:
        issues.append(HealthIssue(
            issue_type="warning",
            category="api",
            title="High Error Rate",
            description=f"Error rate is at {metrics.error_rate:.1f}%",
            suggested_fix="Check API keys and service connectivity",
            auto_fixable=True,
            severity=4 if metrics.error_rate > 25 else 3
        ))
    
    return issues

async def perform_comprehensive_health_check() -> SystemHealthStatus:
    """Perform complete system health check"""
    start_time = time.time()
    
    # Check all components
    db_healthy, db_message = await check_database_health()
    api_healthy, api_message, api_details = await check_api_keys_health()
    deps_healthy, deps_message, missing_deps = await check_dependencies()
    
    # Get system metrics
    metrics = get_system_metrics()
    
    # Calculate active sessions
    try:
        metrics.active_sessions = await db.chat_sessions.count_documents({})
    except:
        metrics.active_sessions = 0
    
    # Calculate response time
    metrics.response_time = (time.time() - start_time) * 1000  # milliseconds
    
    # Determine overall status
    critical_issues = []
    warning_issues = []
    
    # Add component-specific issues
    if not db_healthy:
        critical_issues.append(HealthIssue(
            issue_type="critical",
            category="database",
            title="Database Connection Failed",
            description=db_message,
            suggested_fix="Check MongoDB service and connection string",
            auto_fixable=True,
            severity=5
        ))
    
    if not api_healthy:
        if "No valid API keys" in api_message:
            critical_issues.append(HealthIssue(
                issue_type="critical",
                category="api",
                title="No Valid API Keys",
                description=api_message,
                suggested_fix="Check and update API keys in environment variables",
                auto_fixable=False,
                severity=5
            ))
        else:
            warning_issues.append(HealthIssue(
                issue_type="warning",
                category="api",
                title="Partial API Key Issues",
                description=api_message,
                suggested_fix="Check API key validity and account limits",
                auto_fixable=False,
                severity=3
            ))
    
    if not deps_healthy:
        critical_issues.append(HealthIssue(
            issue_type="critical",
            category="dependency",
            title="Missing Dependencies",
            description=deps_message,
            suggested_fix=f"Install missing packages: pip install {' '.join(missing_deps)}",
            auto_fixable=True,
            severity=4
        ))
    
    # Add performance issues
    performance_issues = await analyze_performance_issues(metrics)
    warning_issues.extend(performance_issues)
    
    # Determine overall status
    if critical_issues:
        overall_status = "critical"
    elif warning_issues:
        overall_status = "warning"
    else:
        overall_status = "healthy"
    
    # Calculate uptime
    uptime = (datetime.utcnow() - health_monitor_data["start_time"]).total_seconds()
    
    all_issues = critical_issues + warning_issues
    
    # Update global health data
    health_monitor_data["issues"] = all_issues
    health_monitor_data["last_health_check"] = datetime.utcnow()
    
    return SystemHealthStatus(
        overall_status=overall_status,
        backend_status="healthy" if db_healthy and deps_healthy else "unhealthy",
        frontend_status="healthy",  # Will be updated by frontend monitoring
        database_status="healthy" if db_healthy else "unhealthy",
        api_status="healthy" if api_healthy else "unhealthy",
        metrics=metrics,
        issues=all_issues,
        uptime=uptime
    )

async def apply_auto_fix(issue: HealthIssue) -> dict:
    """Apply automated fixes based on issue type"""
    
    if issue.category == "dependency":
        # Auto-install missing dependencies
        if "Missing packages:" in issue.description:
            # Extract package names from description
            packages = issue.description.split("Missing packages: ")[1].split(", ")
            install_results = []
            
            for package in packages:
                try:
                    result = subprocess.run([
                        sys.executable, "-m", "pip", "install", package
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0:
                        install_results.append(f"✅ {package} installed successfully")
                    else:
                        install_results.append(f"❌ {package} failed: {result.stderr}")
                except subprocess.TimeoutExpired:
                    install_results.append(f"⏱️ {package} installation timed out")
                except Exception as e:
                    install_results.append(f"❌ {package} error: {str(e)}")
            
            return {
                "action": "install_dependencies",
                "result": "partial" if any("❌" in r for r in install_results) else "success",
                "message": "\n".join(install_results)
            }
    
    elif issue.category == "performance":
        if "High CPU Usage" in issue.title or "High Memory Usage" in issue.title:
            # Clear system caches and suggest restart
            try:
                # Clear Python caches
                import gc
                gc.collect()
                
                # Clear any application caches if they exist
                # (You can add more cache clearing logic here)
                
                return {
                    "action": "clear_caches",
                    "result": "success",
                    "message": "System caches cleared. Consider restarting services if issue persists."
                }
            except Exception as e:
                return {
                    "action": "clear_caches",
                    "result": "failed",
                    "message": f"Cache clearing failed: {str(e)}"
                }
    
    elif issue.category == "service":
        # Service restart suggestions
        return {
            "action": "service_restart_suggestion",
            "result": "info",
            "message": "Consider restarting the backend service using: sudo supervisorctl restart backend"
        }
    
    elif issue.category == "database":
        # Database reconnection attempt
        try:
            global client, db
            client.close()
            client = AsyncIOMotorClient(MONGO_URL)
            db = client[DB_NAME]
            
            # Test connection
            await client.admin.command('ping')
            
            return {
                "action": "database_reconnect",
                "result": "success",
                "message": "Database connection re-established successfully"
            }
        except Exception as e:
            return {
                "action": "database_reconnect",
                "result": "failed",
                "message": f"Database reconnection failed: {str(e)}"
            }
    
    else:
        return {
            "action": "no_fix_available",
            "result": "info",
            "message": "No automated fix available for this issue type"
        }

# Global health monitoring variables
health_monitor_data = {
    "start_time": datetime.utcnow(),
    "api_calls": 0,
    "errors": 0,
    "issues": [],
    "last_health_check": None,
    "metrics_history": []
}

# Health Monitoring Functions
async def check_database_health() -> tuple[bool, str]:
    """Check MongoDB connection health"""
    try:
        # Try to ping the database
        await client.admin.command('ping')
        # Check if our database exists and is accessible
        collections = await db.list_collection_names()
        return True, f"Database healthy - {len(collections)} collections available"
    except Exception as e:
        return False, f"Database connection failed: {str(e)}"

async def check_api_keys_health() -> tuple[bool, str, dict]:
    """Check API keys validity"""
    api_status = {
        "openrouter": {"valid": False, "error": ""},
        "gemini": {"valid": False, "error": ""}
    }
    
    # Check OpenRouter API
    if OPENROUTER_API_KEY:
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(
                    f"{OPENROUTER_BASE_URL}/models",
                    headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}"},
                    timeout=10.0
                )
                if response.status_code == 200:
                    api_status["openrouter"]["valid"] = True
                else:
                    api_status["openrouter"]["error"] = f"HTTP {response.status_code}"
        except Exception as e:
            api_status["openrouter"]["error"] = str(e)
    else:
        api_status["openrouter"]["error"] = "API key not configured"
    
    # Check Gemini API
    if GEMINI_API_KEY:
        try:
            # Test with a simple message
            session_id = str(uuid.uuid4())
            chat = LlmChat(
                api_key=GEMINI_API_KEY,
                session_id=session_id,
                system_message="Test"
            ).with_model("gemini", "gemini-1.5-flash")
            
            # This is a simple test - in production you might want a lighter check
            api_status["gemini"]["valid"] = True
        except Exception as e:
            api_status["gemini"]["error"] = str(e)
    else:
        api_status["gemini"]["error"] = "API key not configured"
    
    # Determine overall API health
    valid_apis = sum(1 for api in api_status.values() if api["valid"])
    total_apis = len([key for key in [OPENROUTER_API_KEY, GEMINI_API_KEY] if key])
    
    if valid_apis == 0:
        return False, "No valid API keys", api_status
    elif valid_apis < total_apis:
        return True, f"{valid_apis}/{total_apis} API keys valid", api_status
    else:
        return True, "All API keys valid", api_status

async def check_dependencies() -> tuple[bool, str, list]:
    """Check if all required dependencies are installed"""
    required_packages = [
        'fastapi', 'uvicorn', 'motor', 'pymongo', 'httpx', 'PyPDF2',
        'anthropic', 'emergentintegrations', 'psutil', 'reportlab',
        'python-docx', 'python-dotenv', 'pydantic'
    ]
    
    missing_packages = []
    
    for package in required_packages:
        try:
            pkg_resources.get_distribution(package)
        except pkg_resources.DistributionNotFound:
            missing_packages.append(package)
    
    if missing_packages:
        return False, f"Missing packages: {', '.join(missing_packages)}", missing_packages
    else:
        return True, "All dependencies installed", []

def get_system_metrics() -> HealthMetrics:
    """Get current system performance metrics"""
    try:
        cpu_usage = psutil.cpu_percent()
        memory = psutil.virtual_memory()
        disk = psutil.disk_usage('/')
        
        # Calculate error rate
        total_calls = health_monitor_data["api_calls"]
        total_errors = health_monitor_data["errors"]
        error_rate = (total_errors / total_calls * 100) if total_calls > 0 else 0
        
        return HealthMetrics(
            cpu_usage=cpu_usage,
            memory_usage=memory.percent,
            disk_usage=disk.percent,
            response_time=0.0,  # Will be calculated during API calls
            active_sessions=0,  # Will be queried from database
            total_api_calls=total_calls,
            error_rate=error_rate
        )
    except Exception as e:
        logger.error(f"Error getting system metrics: {e}")
        return HealthMetrics(
            cpu_usage=0, memory_usage=0, disk_usage=0, response_time=0,
            active_sessions=0, total_api_calls=0, error_rate=0
        )

async def analyze_performance_issues(metrics: HealthMetrics) -> List[HealthIssue]:
    """Analyze metrics and identify performance issues"""
    issues = []
    
    # High CPU usage
    if metrics.cpu_usage > 80:
        issues.append(HealthIssue(
            issue_type="performance",
            category="performance",
            title="High CPU Usage",
            description=f"CPU usage is at {metrics.cpu_usage:.1f}%",
            suggested_fix="Consider restarting services or optimizing queries",
            auto_fixable=True,
            severity=3 if metrics.cpu_usage > 90 else 2
        ))
    
    # High memory usage
    if metrics.memory_usage > 80:
        issues.append(HealthIssue(
            issue_type="performance",
            category="performance",
            title="High Memory Usage",
            description=f"Memory usage is at {metrics.memory_usage:.1f}%",
            suggested_fix="Clear caches and restart services",
            auto_fixable=True,
            severity=3 if metrics.memory_usage > 90 else 2
        ))
    
    # High error rate
    if metrics.error_rate > 10:
        issues.append(HealthIssue(
            issue_type="warning",
            category="api",
            title="High Error Rate",
            description=f"Error rate is at {metrics.error_rate:.1f}%",
            suggested_fix="Check API keys and service connectivity",
            auto_fixable=True,
            severity=4 if metrics.error_rate > 25 else 3
        ))
    
    return issues

async def perform_comprehensive_health_check() -> SystemHealthStatus:
    """Perform complete system health check"""
    start_time = time.time()
    
    # Check all components
    db_healthy, db_message = await check_database_health()
    api_healthy, api_message, api_details = await check_api_keys_health()
    deps_healthy, deps_message, missing_deps = await check_dependencies()
    
    # Get system metrics
    metrics = get_system_metrics()
    
    # Calculate active sessions
    try:
        metrics.active_sessions = await db.chat_sessions.count_documents({})
    except:
        metrics.active_sessions = 0
    
    # Calculate response time
    metrics.response_time = (time.time() - start_time) * 1000  # milliseconds
    
    # Determine overall status
    critical_issues = []
    warning_issues = []
    
    # Add component-specific issues
    if not db_healthy:
        critical_issues.append(HealthIssue(
            issue_type="critical",
            category="database",
            title="Database Connection Failed",
            description=db_message,
            suggested_fix="Check MongoDB service and connection string",
            auto_fixable=True,
            severity=5
        ))
    
    if not api_healthy:
        if "No valid API keys" in api_message:
            critical_issues.append(HealthIssue(
                issue_type="critical",
                category="api",
                title="No Valid API Keys",
                description=api_message,
                suggested_fix="Check and update API keys in environment variables",
                auto_fixable=False,
                severity=5
            ))
        else:
            warning_issues.append(HealthIssue(
                issue_type="warning",
                category="api",
                title="Partial API Key Issues",
                description=api_message,
                suggested_fix="Check API key validity and account limits",
                auto_fixable=False,
                severity=3
            ))
    
    if not deps_healthy:
        critical_issues.append(HealthIssue(
            issue_type="critical",
            category="dependency",
            title="Missing Dependencies",
            description=deps_message,
            suggested_fix=f"Install missing packages: pip install {' '.join(missing_deps)}",
            auto_fixable=True,
            severity=4
        ))
    
    # Add performance issues
    performance_issues = await analyze_performance_issues(metrics)
    warning_issues.extend(performance_issues)
    
    # Determine overall status
    if critical_issues:
        overall_status = "critical"
    elif warning_issues:
        overall_status = "warning"
    else:
        overall_status = "healthy"
    
    # Calculate uptime
    uptime = (datetime.utcnow() - health_monitor_data["start_time"]).total_seconds()
    
    all_issues = critical_issues + warning_issues
    
    # Update global health data
    health_monitor_data["issues"] = all_issues
    health_monitor_data["last_health_check"] = datetime.utcnow()
    
    return SystemHealthStatus(
        overall_status=overall_status,
        backend_status="healthy" if db_healthy and deps_healthy else "unhealthy",
        frontend_status="healthy",  # Will be updated by frontend monitoring
        database_status="healthy" if db_healthy else "unhealthy",
        api_status="healthy" if api_healthy else "unhealthy",
        metrics=metrics,
        issues=all_issues,
        uptime=uptime
    )

# PDF Processing Functions
async def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")

# API Routes
@api_router.post("/sessions", response_model=ChatSession)
async def create_session(request: CreateSessionRequest):
    session = ChatSession(
        title=request.title
    )
    await db.chat_sessions.insert_one(session.dict())
    return session

@api_router.get("/sessions", response_model=List[ChatSession])
async def get_sessions():
    sessions = await db.chat_sessions.find().sort("updated_at", -1).to_list(100)
    return [ChatSession(**session) for session in sessions]

@api_router.post("/sessions/{session_id}/upload-pdf")
async def upload_pdf(session_id: str, file: UploadFile = File(...)):
    # Verify session exists
    session = await db.chat_sessions.find_one({"id": session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    # Read and process PDF
    file_content = await file.read()
    pdf_text = await extract_text_from_pdf(file_content)
    
    # Save PDF document
    pdf_doc = PDFDocument(
        filename=file.filename,
        content=pdf_text,
        file_size=len(file_content)
    )
    await db.pdf_documents.insert_one(pdf_doc.dict())
    
    # Update session with PDF info
    await db.chat_sessions.update_one(
        {"id": session_id},
        {
            "$set": {
                "pdf_filename": file.filename,
                "pdf_content": pdf_text,
                "updated_at": datetime.utcnow()
            }
        }
    )
    
    return {
        "message": "PDF uploaded successfully",
        "filename": file.filename,
        "content_length": len(pdf_text)
    }

@api_router.post("/sessions/{session_id}/messages")
async def send_message(session_id: str, request: SendMessageRequest):
    # Verify session exists
    session = await db.chat_sessions.find_one({"id": session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Save user message
    user_message = ChatMessage(
        session_id=session_id,
        content=request.content,
        role="user",
        feature_type=request.feature_type
    )
    await db.chat_messages.insert_one(user_message.dict())
    
    # Get chat history
    messages_cursor = db.chat_messages.find({"session_id": session_id}).sort("timestamp", 1)
    chat_history = await messages_cursor.to_list(100)
    
    # Prepare messages for AI based on feature type
    ai_messages = []
    
    if request.feature_type == "general_ai":
        # General AI chat - no PDF context
        ai_messages.append({
            "role": "system", 
            "content": "You are a helpful AI assistant. Answer any questions the user has with accurate and helpful information."
        })
    else:
        # PDF-based features
        if session.get("pdf_content"):
            pdf_content = session["pdf_content"]
            if request.feature_type == "chat":
                system_message = f"""You are an AI assistant specialized in analyzing PDF documents. 

PDF Content:
{pdf_content[:4000]}...

Please answer questions based on this PDF content. Be specific and reference the document when possible."""
            
            ai_messages.append({"role": "system", "content": system_message})
        else:
            ai_messages.append({
                "role": "system", 
                "content": "You are a helpful AI assistant. No PDF has been uploaded yet. Please ask the user to upload a PDF document first."
            })
    
    # Add recent conversation history
    for msg in chat_history[-10:]:  # Last 10 messages for context
        if msg["feature_type"] == request.feature_type or request.feature_type == "chat":
            ai_messages.append({
                "role": msg["role"],
                "content": msg["content"]
            })
    
    # Get AI response
    ai_response = await get_ai_response(ai_messages, request.model)
    
    # Save AI message
    ai_message = ChatMessage(
        session_id=session_id,
        content=ai_response,
        role="assistant",
        feature_type=request.feature_type
    )
    await db.chat_messages.insert_one(ai_message.dict())
    
    # Update session timestamp
    await db.chat_sessions.update_one(
        {"id": session_id},
        {"$set": {"updated_at": datetime.utcnow()}}
    )
    
    return {"ai_response": ai_message}

@api_router.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    # Delete session
    result = await db.chat_sessions.delete_one({"id": session_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Delete associated messages
    await db.chat_messages.delete_many({"session_id": session_id})
    
    return {"message": "Session deleted successfully"}

@api_router.get("/sessions/{session_id}/messages", response_model=List[ChatMessage])
async def get_messages(session_id: str, feature_type: Optional[str] = Query(None)):
    # Verify session exists
    session = await db.chat_sessions.find_one({"id": session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Build query
    query = {"session_id": session_id}
    if feature_type:
        query["feature_type"] = feature_type
    
    messages = await db.chat_messages.find(query).sort("timestamp", 1).to_list(1000)
    return [ChatMessage(**message) for message in messages]

@api_router.get("/models")
async def get_available_models():
    models = []
    
    # Add OpenRouter models (Claude) if API key is configured
    if OPENROUTER_API_KEY:
        models.extend([
            {
                "id": "claude-3-opus-20240229",
                "name": "Claude 3 Opus",
                "provider": "OpenRouter",
                "free": False
            },
            {
                "id": "claude-3-sonnet-20240229",
                "name": "Claude 3 Sonnet",
                "provider": "OpenRouter",
                "free": False
            },
            {
                "id": "claude-3-haiku-20240307",
                "name": "Claude 3 Haiku",
                "provider": "OpenRouter",
                "free": False
            }
        ])
    
    # Add Gemini models if API key is configured
    if GEMINI_API_KEY:
        models.extend([
            {
                "id": "gemini-2.0-flash",
                "name": "Gemini 2.0 Flash",
                "provider": "Google",
                "free": True
            },
            {
                "id": "gemini-1.5-flash",
                "name": "Gemini 1.5 Flash",
                "provider": "Google",
                "free": True
            },
            {
                "id": "gemini-1.5-pro",
                "name": "Gemini 1.5 Pro",
                "provider": "Google",
                "free": True
            },
            {
                "id": "gemini-1.5-flash-8b",
                "name": "Gemini 1.5 Flash 8B",
                "provider": "Google",
                "free": True
            }
        ])
    
    return {"models": models}



# Removed generate-qa endpoint - replaced with generate-questions

# Removed research endpoint - replaced with new features

@api_router.post("/translate")
async def translate_pdf(request: TranslateRequest):
    # Verify session exists and has PDF
    session = await db.chat_sessions.find_one({"id": request.session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if not session.get("pdf_content"):
        raise HTTPException(status_code=400, detail="No PDF uploaded in this session")
    
    pdf_content = session["pdf_content"]
    
    # Limit content based on type
    if request.content_type == "summary":
        content_to_translate = pdf_content[:2000]
        translation_instruction = f"Provide a translated summary of this document in {request.target_language}:"
    else:
        content_to_translate = pdf_content[:4000]
        translation_instruction = f"Translate this document content to {request.target_language}:"
    
    ai_messages = [
        {
            "role": "system", 
            "content": f"You are a professional translator. Translate the given content accurately to {request.target_language} while maintaining the original meaning and context."
        },
        {
            "role": "user", 
            "content": f"""{translation_instruction}

{content_to_translate}"""
        }
    ]
    
    translation_result = await get_ai_response(ai_messages, request.model)
    
    # Save translation as message
    translation_message = ChatMessage(
        session_id=request.session_id,
        content=f"Translation to {request.target_language} ({request.content_type}):\n{translation_result}",
        role="assistant",
        feature_type="translation"
    )
    await db.chat_messages.insert_one(translation_message.dict())
    
    return {
        "session_id": request.session_id,
        "target_language": request.target_language,
        "content_type": request.content_type,
        "translation": translation_result
    }

@api_router.post("/generate-questions")
async def generate_questions(request: GenerateQuestionsRequest):
    # Verify session exists and has PDF
    session = await db.chat_sessions.find_one({"id": request.session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if not session.get("pdf_content"):
        raise HTTPException(status_code=400, detail="No PDF uploaded in this session")
    
    pdf_content = session["pdf_content"]
    
    # Handle chapter segmentation if specified
    if request.chapter_segment:
        # Simple chapter detection - could be enhanced
        content_lines = pdf_content.split('\n')
        chapter_content = []
        in_chapter = False
        
        for line in content_lines:
            if request.chapter_segment.lower() in line.lower():
                in_chapter = True
            elif any(keyword in line.lower() for keyword in ['chapter', 'section']) and in_chapter:
                break
            
            if in_chapter:
                chapter_content.append(line)
        
        pdf_content = '\n'.join(chapter_content)[:4000] if chapter_content else pdf_content[:4000]
    else:
        pdf_content = pdf_content[:4000]
    
    # Question generation prompts based on type
    question_prompts = {
        "faq": "Generate 8-10 frequently asked questions (FAQs) with detailed answers based on this document content.",
        "mcq": "Generate 10 multiple choice questions (A, B, C, D) with correct answers marked, based on this document content.",
        "true_false": "Generate 10 true/false questions with explanations for each answer, based on this document content.",
        "mixed": "Generate a mix of question types: 3 FAQs, 4 multiple choice questions, and 3 true/false questions based on this document content."
    }
    
    prompt = question_prompts.get(request.question_type, question_prompts["mixed"])
    
    ai_messages = [
        {
            "role": "system", 
            "content": "You are an AI assistant specialized in creating educational questions from document content. Generate clear, relevant questions that test comprehension and knowledge retention."
        },
        {
            "role": "user", 
            "content": f"""{prompt}

Document Content:
{pdf_content}

Format your response clearly with question numbers, and for MCQs include all options (A, B, C, D) with the correct answer marked."""
        }
    ]
    
    questions_result = await get_ai_response(ai_messages, request.model)
    
    # Save questions as message
    questions_message = ChatMessage(
        session_id=request.session_id,
        content=f"Generated Questions ({request.question_type}):\n{questions_result}",
        role="assistant",
        feature_type="question_generation"
    )
    await db.chat_messages.insert_one(questions_message.dict())
    
    return {
        "session_id": request.session_id,
        "question_type": request.question_type,
        "chapter_segment": request.chapter_segment,
        "questions": questions_result
    }

@api_router.post("/generate-quiz")
async def generate_quiz(request: GenerateQuizRequest):
    # Verify session exists and has PDF
    session = await db.chat_sessions.find_one({"id": request.session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if not session.get("pdf_content"):
        raise HTTPException(status_code=400, detail="No PDF uploaded in this session")
    
    pdf_content = session["pdf_content"][:4000]  # Limit content length
    
    # Difficulty level instructions
    difficulty_instructions = {
        "easy": "Create basic comprehension questions that test understanding of main concepts.",
        "medium": "Create questions that require analysis and application of the content.",
        "hard": "Create challenging questions that require deep understanding, critical thinking, and synthesis of multiple concepts."
    }
    
    difficulty_instruction = difficulty_instructions.get(request.difficulty, difficulty_instructions["medium"])
    
    # Quiz type specific instructions
    if request.quiz_type == "daily":
        quiz_instruction = f"Generate a daily revision quiz with {request.question_count} questions. Focus on key concepts for daily review."
    else:
        quiz_instruction = f"Generate a comprehensive quiz with {request.question_count} questions covering the document content."
    
    ai_messages = [
        {
            "role": "system", 
            "content": f"You are an AI quiz generator specialized in creating educational quizzes. {difficulty_instruction} Make questions clear and provide correct answers."
        },
        {
            "role": "user", 
            "content": f"""{quiz_instruction}

Document Content:
{pdf_content}

Create {request.question_count} questions in mixed format (multiple choice, true/false, short answer). 
For each question, provide:
1. The question
2. Answer options (if applicable)
3. Correct answer
4. Brief explanation

Format the quiz clearly with question numbers."""
        }
    ]
    
    quiz_result = await get_ai_response(ai_messages, request.model)
    
    # Save quiz as message
    quiz_message = ChatMessage(
        session_id=request.session_id,
        content=f"Generated Quiz ({request.quiz_type} - {request.difficulty}):\n{quiz_result}",
        role="assistant",
        feature_type="quiz_generation"
    )
    await db.chat_messages.insert_one(quiz_message.dict())
    
    return {
        "session_id": request.session_id,
        "quiz_type": request.quiz_type,
        "difficulty": request.difficulty,
        "question_count": request.question_count,
        "quiz": quiz_result
    }

@api_router.post("/search")
async def advanced_search(request: SearchRequest):
    results = []
    
    if request.search_type in ["all", "pdfs"]:
        # Search in PDF documents
        pdf_query = {"content": {"$regex": request.query, "$options": "i"}}
        pdf_docs = await db.pdf_documents.find(pdf_query).limit(request.limit).to_list(request.limit)
        
        for doc in pdf_docs:
            # Find snippet around the search term
            content = doc["content"]
            query_lower = request.query.lower()
            content_lower = content.lower()
            
            if query_lower in content_lower:
                start_idx = max(0, content_lower.find(query_lower) - 100)
                end_idx = min(len(content), start_idx + 300)
                snippet = content[start_idx:end_idx]
                
                results.append({
                    "type": "pdf",
                    "filename": doc["filename"],
                    "snippet": snippet,
                    "upload_date": doc["upload_date"],
                    "relevance_score": content_lower.count(query_lower)
                })
    
    if request.search_type in ["all", "conversations"]:
        # Search in chat messages
        msg_query = {"content": {"$regex": request.query, "$options": "i"}}
        messages = await db.chat_messages.find(msg_query).limit(request.limit).to_list(request.limit)
        
        for msg in messages:
            # Get session info for context
            session = await db.chat_sessions.find_one({"id": msg["session_id"]})
            session_title = session["title"] if session else "Unknown Session"
            
            results.append({
                "type": "conversation",
                "session_title": session_title,
                "session_id": msg["session_id"],
                "content": msg["content"][:200] + "..." if len(msg["content"]) > 200 else msg["content"],
                "timestamp": msg["timestamp"],
                "feature_type": msg["feature_type"]
            })
    
    # Sort results by relevance/recency
    results.sort(key=lambda x: x.get("relevance_score", 1) if x["type"] == "pdf" else 1, reverse=True)
    
    return {
        "query": request.query,
        "search_type": request.search_type,
        "total_results": len(results),
        "results": results[:request.limit]
    }

@api_router.post("/export")
async def export_conversation(request: ExportRequest):
    # Verify session exists
    session = await db.chat_sessions.find_one({"id": request.session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Get messages based on filter
    query = {"session_id": request.session_id}
    if request.feature_type:
        query["feature_type"] = request.feature_type
    
    messages = await db.chat_messages.find(query).sort("timestamp", 1).to_list(1000)
    
    # Format content for export
    export_content = f"Chat Session: {session['title']}\n"
    export_content += f"Created: {session['created_at']}\n"
    if session.get('pdf_filename'):
        export_content += f"PDF: {session['pdf_filename']}\n"
    export_content += "\n" + "="*50 + "\n\n"
    
    for msg in messages:
        role_label = "User" if msg["role"] == "user" else "Assistant"
        export_content += f"[{msg['timestamp']}] {role_label}:\n{msg['content']}\n\n"
    
    if request.export_format == "txt":
        # Return plain text
        from fastapi.responses import Response
        return Response(
            content=export_content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename={session['title']}.txt"}
        )
    elif request.export_format == "pdf":
        # Generate PDF using reportlab
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from io import BytesIO
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title = Paragraph(f"Chat Session: {session['title']}", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Add messages
        for msg in messages:
            role_label = "User" if msg["role"] == "user" else "Assistant"
            header = Paragraph(f"{role_label} ({msg['timestamp']}):", styles['Heading2'])
            content = Paragraph(msg['content'], styles['Normal'])
            story.append(header)
            story.append(content)
            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        
        from fastapi.responses import Response
        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={session['title']}.pdf"}
        )
    elif request.export_format == "docx":
        # Generate DOCX using python-docx
        from docx import Document
        from io import BytesIO
        
        doc = Document()
        doc.add_heading(f"Chat Session: {session['title']}", 0)
        
        # Add session info
        doc.add_paragraph(f"Created: {session['created_at']}")
        if session.get('pdf_filename'):
            doc.add_paragraph(f"PDF: {session['pdf_filename']}")
        
        # Add messages
        for msg in messages:
            role_label = "User" if msg["role"] == "user" else "Assistant"
            doc.add_heading(f"{role_label} ({msg['timestamp']})", level=2)
            doc.add_paragraph(msg['content'])
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from fastapi.responses import Response
        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={session['title']}.docx"}
        )
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported export format")

@api_router.get("/insights")
async def get_insights():
    # Get total sessions
    total_sessions = await db.chat_sessions.count_documents({})
    
    # Get total messages
    total_messages = await db.chat_messages.count_documents({})
    
    # Get total PDFs uploaded
    total_pdfs = await db.pdf_documents.count_documents({})
    
    # Get feature usage statistics
    feature_usage = await db.chat_messages.aggregate([
        {"$group": {"_id": "$feature_type", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]).to_list(10)
    
    # Get popular PDFs (most referenced)
    popular_pdfs = await db.chat_sessions.aggregate([
        {"$match": {"pdf_filename": {"$exists": True, "$ne": None}}},
        {"$group": {"_id": "$pdf_filename", "usage_count": {"$sum": 1}}},
        {"$sort": {"usage_count": -1}},
        {"$limit": 5}
    ]).to_list(5)
    
    # Get daily usage for last 7 days
    from datetime import datetime, timedelta
    seven_days_ago = datetime.utcnow() - timedelta(days=7)
    
    daily_usage = await db.chat_messages.aggregate([
        {"$match": {"timestamp": {"$gte": seven_days_ago}}},
        {
            "$group": {
                "_id": {
                    "$dateToString": {
                        "format": "%Y-%m-%d",
                        "date": "$timestamp"
                    }
                },
                "message_count": {"$sum": 1}
            }
        },
        {"$sort": {"_id": 1}}
    ]).to_list(7)
    
    return {
        "overview": {
            "total_sessions": total_sessions,
            "total_messages": total_messages,
            "total_pdfs": total_pdfs
        },
        "feature_usage": feature_usage,
        "popular_pdfs": popular_pdfs,
        "daily_usage": daily_usage
    }

# Add CORS middleware with environment-specific origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS if ENVIRONMENT == 'production' else ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include API routes
app.include_router(api_router)